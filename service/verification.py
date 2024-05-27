import hashlib
import requests
from docx import Document
import re

import zipfilefrom
from lxml import etree

# Отделение ключа от документа, для последующей проверки
def check_docx_signature(docx_path):
    with zipfile.ZipFile(docx_path, 'r') as docx:
        if 'docProps/core.xml' in docx.namelist():
            core_properties = docx.read('docProps/core.xml')
            root = etree.fromstring(core_properties)
            namespaces = {'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties'}
            for signature in root.findall('.//cp:keywords', namespaces):
                if 'signed' in signature.text.lower():
                    return True
    return False
    docx_path = 'path/to/document.docx'
    if check_docx_signature(docx_path):
        print("Документ содержит электронную подпись.")
    else:
        print("Документ не содержит электронную подпись.")

def compare_docx_contents(file_path1, file_path2):
    doc1 = Document(file_path1)

    text1 = ''
    for paragraph in doc1.paragraphs:
        text1 += paragraph.text

    doc2 = Document(file_path2)

    text2 = ''
    for paragraph in doc2.paragraphs:
        text2 += paragraph.text

    return text1 == text2


def main():
    token = request.json.get('token')
    original_file_path = '/home/oleg/PycharmProjects/pythonProject1/helper/document0yh55nhmO4uB5BELXCKl65sktq1JUbJjP0MCgWoX.docx'
    returned_file_path = '/home/oleg/PycharmProjects/pythonProject1/helper/document0yh55nhmO4uB5BELXCKl65sktq1JUbJjP0MCgWoX.docx'

    if not compare_docx_contents(original_file_path,returned_file_path):
        return jsonify({'error': 'Файл был изменён. Документ не валиден'}), 401

    return jsonify({'error': 'Файл не был изменён. Документ валиден'}), 200


if __name__ == "__main__":
    main()
