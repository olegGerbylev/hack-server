from flask import request, jsonify, g
import psycopg2
import psycopg2.extras
import validation.main as validation
from docx import Document
from datetime import datetime, timedelta
from service.sender import send_email

def replace_multiple_text_in_docx(docx_file, replacements, token):
    doc = Document(docx_file)

    for old_text, new_text in replacements.items():
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)

    doc.save(f"/home/oleg/PycharmProjects/pythonProject1/document/document{token}.docx")


def sendDocx():
    db = g.db
    cursor = db.cursor(
        cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        email = request.json.get('email')
        validation.validateEmail(email)
        token = request.json.get('token')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    cursor.execute('''SELECT * FROM users WHERE token = %s ''', (token,))
    user = cursor.fetchone()

    if user is None:
        return jsonify({'error': 'Ошибка авторизации'}), 401

    replacements = {
        "%компания": user['company'],
        "%ИНН_компании": user['company_tin'],
        "%ФИО": f'{user["name"]} {user["family_name"]} {user["surname"]}',
        "%серия_паспорта": user['passport_series'],
        "%номер_паспорта": user['passport_id'],
        "%кем_выдан": user['issued_by'],
        "%код_подразделения": user['department_code'],
        "%когда_выдан": user['date_issue'],
        "%адрес_регистрации": user['location'],
        "%дата_окончания_договора": (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d")
    }

    replace_multiple_text_in_docx('/home/oleg/PycharmProjects/pythonProject1/document/original.docx', replacements, token)

    send_email(f"/home/oleg/PycharmProjects/pythonProject1/document/document{token}.docx", email)

    return jsonify({'message': 'Файл успешно отправлен'})